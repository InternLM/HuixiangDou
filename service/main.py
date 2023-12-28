try:
    import thread
except ImportError:
    import _thread as thread
import time
import json

from model_qwen_client import QwenClient
from feature_store import FeatureStore
from loguru import logger
import argparse
import os
import pdb
import requests
import wx_login
import security_remote
import random
from web_search import Searcher
from source_graph_search import sg_search
import datetime
import re
import aliyunoss
from docx import Document

# ocrl handler

# about LLM
parser = argparse.ArgumentParser()
parser.add_argument("--feature_dir", "-f", type=str, default="memory", help="path feature store feature dir")
parser.add_argument("--device_id", "-d", type=int, default=7, help="GPU device id for sentence_transformer")

args, _ = parser.parse_known_args()
device = 'cuda:2'
fs = FeatureStore(device=device)
fs.load_feature()
MAX_INPUT_LEN = 14000
llm = QwenClient()

# about google
google = Searcher()

# about server
remote_ip_port = '101.133.161.204:8888'
me = None
with open('res/me.txt') as f:
    me = json.load(f)

# 已发送的消息列表
sent_msg = dict()

# 所有人的消息列表
all_msgs = dict()



def single_judge(prompt, throttle: int, default: int, repeat=1):
    scores = []

    for _ in range(repeat):
        score = default
        relation = llm.generate_response(prompt=prompt, remote=False)
        filtered_relation = ''.join([c for c in relation if c.isdigit()])
        try:
            score_str = re.sub(r'[^\d]', ' ', filtered_relation).strip()
            score = int(score_str.split(' ')[0])
        except:
            pass

        scores.append(score)

    avg_score = sum(scores) / repeat
    if avg_score >= throttle:
        return True
    return False



def generate(question_simple, question_full, history, groupname):
    log_tracker = dict()
    log_tracker['step1_is_question'] = [question_full, history, groupname]

    logger.info('question full:{}'.format(question_full))
    
    response = ''
    send = False
    reason = ''

    prompt = '''判断以下句子是否是个有主题的疑问句，结果用 1～10 表示。直接提供得分不要解释。
判断标准：有主语谓语宾语并且是疑问句得 10 分；缺少主谓宾扣分；陈述句直接得 0 分；不是疑问句直接得 0 分。

问题“{}”，得分是多少？直接提供得分不要解释。'''.format(question_full)

    # prompt = '''
    # 判断以下句子是否是个有主题的疑问句，结果用 1～10 表示。直接提供得分不要解释。
    # 判断标准：有主语谓语宾语应该得 10 分；缺少关键要素扣分；陈述句直接得 0 分；不是疑问句直接得 0 分。
    # 问题“{}”，得分是多少？直接提供得分不要解释。
    # '''.format(question_full)

    if not single_judge(prompt=prompt, throttle=7, default=3):
        reason = 'not a question'
        log_tracker['step1_not_question'] = 'true'
        return response, send, reason
    
    topic = llm.generate_response('告诉我这句话的主题，直接说主题不要解释：“{}”'.format(question_full))

    log_tracker['step2_topic'] = topic

    if len(topic) <= 2:
        return response, send, reason
    
    db_context_part, db_context = fs.query_source(topic)

    if db_context != '<reject>':

        if single_judge('判断问题和材料的关联度，用1～10表示。判断标准：非常相关得 10 分；完全没关联得 0 分。直接提供得分不要解释。\n问题：“{}”\n材料：“{}”'.format(question_full, db_context_part), throttle=5, default=10):
            prompt, history = llm.build_prompt(instruction=question_full, context=db_context, history_pair=history)
            response = llm.generate_response(prompt=prompt, history=history, remote=True)
            logger.info('-question:{}\n-dbcontext:{}\n-resp:{}\n'.format(question_full, db_context, response))
            log_tracker['step3_1_doc'] = [db_context_part, response]

        else:
            prompt = '''以下是一些名词解释：
谷歌搜索是一个通用搜索引擎，可用于访问互联网、查询百科知识、了解时事新闻等。搜索参数类型 string， 内容是短语或关键字，以空格分隔。 
OpenCompass 是一个开源开放的大模型评测平台，构建了包含学科、语言、知识、理解、推理五大维度的通用能力评测体系，支持了超过 50 个评测数据集和 30 万道评测题目，支持零样本、小样本及思维链评测，是目前最全面的开源评测平台。
XTuner 是低成本大模型训练工具箱，旨在让大模型训练不再有门槛。通过 XTuner，最低只需 8GB 显存，就可以打造专属于你的 AI 助手。
lmdeploy 是涵盖了 LLM 任务的全套轻量化、部署和服务解决方案。基于 FasterTransformer，lmdeploy 实现了高效推理引擎 TurboMind，支持 InternLM、LLaMA、vicuna等模型在 NVIDIA GPU 上的推理。通过缓存多轮对话过程中 attention 的 k/v，记住对话历史，从而避免重复处理历史会话。提供了全面的模型部署和量化支持，已在不同规模上完成验证。
“硬件模型库” 是 openmmlab 部署的一个 web 服务，包含 2500 个已经转换好的视觉模型、nvidia 全套硬件。用户可以直接下载模型使用，也可以用硬件完成真实的模型测速。硬件模型库还提供嵌入式芯片测评服务，目前已发布 ax620a 测评报告。

以下是任务描述：
你现在是{}技术交流群里的技术助手，用户问“{}”，你打算通过谷歌搜索查询相关资料，请提供用于搜索的关键字或短语，不要解释直接给出关键字或短语。'''.format(groupname, question_full)
            web_keywords = llm.generate_response(prompt=prompt)
            # format keywords
            web_keywords = web_keywords.replace('"', ' ')
            web_keywords = web_keywords.replace(',', ' ')
            web_keywords = web_keywords.replace('  ', ' ')
            web_keywords = web_keywords.strip()

            log_tracker['step3_2_web_keywords'] = web_keywords


            if web_keywords is not None and len(web_keywords) > 0:
                try:
                    web_context = ''
                    articles = google.get_with_cache(query=web_keywords, max_article=3)
                    logger.debug('search with cache finished')
                    for article in articles:
                        if article is not None and len(article) > 0:
                            if len(article) > MAX_INPUT_LEN:
                                article = article[0:MAX_INPUT_LEN]

                            if single_judge('判断以下句子是否涉及政治、辱骂、色情、恐暴、宗教、网络暴力、种族歧视等违禁内容，结果用 1～10 表示，不要解释直接给出得分。判断标准：涉其中任一问题直接得 10 分；完全不涉及得 0 分。直接给得分不要解释：“{}”'.format(article), throttle=3, default=0):
                                logger.debug('**跳过不安全的内容** {}'.format(article))
                                continue

                            if single_judge('问题：“{}”\n材料：“{}\n请阅读以上内容，判断问题和材料的关联度，用1～10表示，不要解释直接给出得分。判断标准：非常相关得 10 分；完全没关联得 0 分。直接打分不要解释。\n'.format(question_full, article), throttle=6, default=10):
                                web_context += '\n\n'
                                web_context += article

                    if len(web_context) >= MAX_INPUT_LEN:
                        web_context = web_context[0:MAX_INPUT_LEN]

                    web_context = web_context.strip()

                    if len(web_context) > 0:
                        prompt, history = llm.build_prompt(instruction=question_full, context=web_context, history_pair=history)
                        response = llm.generate_response(prompt=prompt, history=history, remote=False)
                        log_tracker['step3_3_web_context'] = [web_context, response]
                    else:
                        reason = 'no web context'
                    logger.info('-question:{}\n-webcontext:{}\n-resp:{}\n'.format(question_full, web_context, response))
                except Exception as e:
                    logger.error(e)

        if response is not None and len(response) > 0:
            prompt = '''“question:{} answer:{}”

仔细阅读以上对话，answer 是否在表达自己不知道，回答越全面得分越少，用0～10表示，不要解释直接给出得分。
判断标准：准确回答问题得 0 分；答案详尽得 1 分；知道部分答案但有不确定信息得 8 分；知道小部分答案但推荐求助其他人得 9 分；不知道任何答案直接推荐求助别人得 10 分。

以下是一些 10 分示例：
answer: “从您提供的信息来看，这个问题似乎是关于某个技术问题的讨论，可能是在寻求帮助。然而，问题的具体内容并不清楚。为了给您提供更有针对性的帮助，能否请您详细描述一下您所遇到的问题或需要解答的疑问？这样我才能更好地理解您的需求并提供相应的解答。”
answer: “抱歉我不清楚 lmdeploy 是什么，请联系相关开发人员”
直接打分不要解释。'''.format(question_full, response)
            if single_judge(prompt=prompt, throttle=8, default=0):
                send = False
                reason = 'bad answer'
                log_tracker['step3_4_web_search_unknown'] = 'true'
            else:
                send = True

        if reason == 'bad answer' or reason == 'no web context':
            # 找不到答案，或者回答的不知道，尝试检索源码
            if groupname in ['lmdeploy', 'opencompass', 'xtuner']:
                # use page indexing
                sg_context = sg_search(llm=llm, question=question_full, groupname=groupname)
                if sg_context != None and len(sg_context) > 0:
                    if len(sg_context) > 60000:
                        sg_context = sg_context[0:60000]
                    prompt, history = llm.build_prompt(instruction=question_full, context=sg_context, history_pair=history)

                    if len(prompt) > MAX_INPUT_LEN:
                        response = llm.generate_response(prompt=prompt, history=history, remote=True)
                        log_tracker['step4_1_source_graph'] = [sg_context, response, 'remote']

                    else:
                        response = llm.generate_response(prompt=prompt, history=history, remote=False)
                        log_tracker['step4_1_source_graph'] = [sg_context, response, 'local']

                    prompt = '''
“question:{} answer:{}”

仔细阅读以上对话，answer 是否在表达自己不知道，用0～10表示，不要解释直接给出得分。
判断标准：准确回答问题得 0 分；答案详尽得 1 分；知道部分答案但有不确定信息得 8 分；知道小部分答案但推荐求助其他人得 9 分；不知道任何答案直接推荐求助别人得 10 分。

以下是一些 10 分示例：
answer: “从您提供的信息来看，这个问题似乎是关于某个技术问题的讨论，可能是在寻求帮助。然而，问题的具体内容并不清楚。为了给您提供更有针对性的帮助，能否请您详细描述一下您所遇到的问题或需要解答的疑问？这样我才能更好地理解您的需求并提供相应的解答。”
answer: “抱歉我不清楚 lmdeploy 是什么，请联系相关开发人员”
直接打分不要解释。
                    '''.format(question_full, response)
                    if single_judge(prompt=prompt, throttle=9, default=0):
                        send = False
                        reason = 'bad answer'
                        log_tracker['step4_2_source_graph_unknown'] = 'true'
                    else:
                        send = True

        if response is not None and len(response) >= 600:
            # 回复内容太长，总结一下
            response = llm.generate_response(prompt='{} \n 仔细阅读以上内容，总结到 500 字以内'.format(response), history=[])

        if send:
            if single_judge('判断以下句子是否涉及政治、辱骂、色情、恐暴、宗教、网络暴力、种族歧视等违禁内容，结果用 1～10 表示，不要解释直接给出得分。判断标准：涉其中任一问题直接得 10 分；完全不涉及得 0 分。直接给得分不要解释：“{}”'.format(response), throttle=3, default=0):
                response = '“茴香豆”作为 LLM 助手，须遵循社会主义核心价值观。'
                send = False
                reason = 'security'
        if send:
            if not security_remote.check(question_simple + '。' + response):
                response = '“茴香豆”作为业务助手，须遵循社会主义核心价值观。'
                send = False
                reason = 'security'

    if response is None:
        response = ''
    log_tracker['step5_final'] = [response, send, reason]
    with open('res/log_tracker.txt', 'a', encoding='utf-8') as f:
        f.write(json.dumps(log_tracker, ensure_ascii=False, indent=2))
        f.write('\n')
        f.write('\n')
        f.write('\n')
    return response, send, reason


def record_history(x):
    with open('res/history_recv_send.txt', 'a') as f:
        jsonstr = json.dumps(x, ensure_ascii=False, indent=2)
        f.write(jsonstr)
        f.write('\n')


def count_history(history: list):
    count = 0
    for item in history:
        count += len(item[0])
        count += len(item[1])
    return count

def save_sent(groupId, question_full, answer, filepath):
    with open(filepath, 'a', encoding='utf-8') as f:
        obj = {'groupId': groupId, 'question_full': question_full, 'answer': answer}
        f.write(json.dumps(obj, ensure_ascii=False, indent=2))
        f.write('\n')

def send_message(groupId, answer: str):
    wId = me['wId']
    success, sent = wx_login.send(wId=wId, group=groupId, content=answer)
    if success:
        record_history({'groupId': groupId, 'answer': answer})
        logger.debug(sent)
        """
        发送成功
        '{"code":"1000","message":"abc","data":{"type":1,"msgId":2193854338,"newMsgId":6863876859054015532,"createTime":1697343933,"wcId":"20158567857@chatroom"}}
        """
        sent = json.loads(sent)['data']
        sent['wId'] = wId
        if groupId not in sent_msg:
            sent_msg[groupId] = [sent]
        else:
            sent_msg[groupId].append(sent)


def revert_all(groupId: str):
    # 撤回在本群 2 分钟内发出的所有消息
    logger.debug('revert message')
    if groupId in sent_msg:
        group_sent_list = sent_msg[groupId]
        for sent in group_sent_list:
            logger.info(sent)
            time_diff = abs(time.time() - int(sent['createTime']))
            if time_diff <= 120:
                wx_login.revert(sent)
                time.sleep(0.5)
        del sent_msg[groupId]


def llm_response(wx_msg, treat_as_context=False):
    data = wx_msg['data']
    # 提取问题内容
    question = data['content']
    question = question.encode('UTF-8', 'ignore').decode('UTF-8')
    fromUser = data['fromUser']
    groupId = data['fromGroup']

    if '@茴香豆' in question and '撤回' in question:
        revert_all(groupId)
        return None

    if '豆哥撤回' in question:
        revert_all(groupId)
        return None

    # PC 版微信回复别人，不处理
    if '————————' in question:
        return None

    # 合并问题
    # key 是 chatroom+fromUser
    # value 是 {"last_time": UTC, history:[(question, response)]}
    # 每收到一个消息, 更新 last_time
    #   如果 8s 内又发了消息，和上一条 merge 成一个 question

    # 检查所有消息
    #   如果 last_time 是 8hour 前的，删除这个人
    #   如果总长度超过 2048，保留最后一条其他删掉
    #   如果超过 15s 且回复时间早于接收时间，开始回复
    #       回复完，更新回复时间

    def update_user():
        # 更新发送方
        now = time.time()
        key = '{}|{}'.format(fromUser, groupId)
        if key not in all_msgs:
            all_msgs[key] = {'history': [(question, '')], 'last_wx_msg_time': now, 'last_response_time': -1, 'group_id': groupId}
            return None
        user = all_msgs[key]
        user['last_wx_msg_time'] = now
        user['history'].append((question, ''))

    update_user()
    if treat_as_context:
        # for context, only update to user message list, do not process
        return None

    def format_history(history: list):
        # 整理历史消息，把没有回复的消息合并
        if len(history) < 2:
            return history
        ret = []
        merge_list = []
        for item in history:
            answer = item[1]
            if answer is not None and len(answer) > 0:
                ret.append(item)
            else:
                merge_list.append(item[0])
        ret.append(('。'.join(merge_list), ''))
        return ret

    def scan_all():
        # 扫描所有人的消息
        for key in all_msgs.keys():
            user = all_msgs[key]
            last_wx_msg_time = user['last_wx_msg_time']
            last_response_time = user['last_response_time']
            # if now - last_wx_msg_time >= 3600 * 4:
            #     del all_msgs[key]
            #     continue
            if count_history(user['history']) >= 10000:
                user['history'] = [user['history'][-1]]
                continue

            now = time.time()
            if now - last_wx_msg_time >= 18 and last_response_time < last_wx_msg_time:
                history = user['history']

                question_simple = history[-1][0]
                history = format_history(history=history)

                question_full = history[-1][0]

                groupname = 'openmmlab'
                groupid2name = {
                    '21375247392@chatroom': 'opencompass',
                    '35003485913@chatroom': 'opencompass',
                    '20158567857@chatroom': 'opencompass',
                    '21295744750@chatroom': '硬件模型库',
                    '34744063953@chatroom': 'xtuner',
                    '21177113665@chatroom': 'lmdeploy',
                    '20814553575@chatroom': 'openmmlab',
                    '18356748488@chatroom': 'ncnn'
                }
                groupId = user['group_id']
                if groupId in groupid2name:
                    groupname = groupid2name[groupId]

                answer = ''
                send = False
                reason = ''
                # 简单词汇直接当成上下文塞进 context。响应的最简单的问题 “如何安装mmdet”
                if len(question_simple) >= 9:
                    answer, send, reason = generate(question_simple=question_simple, question_full=question_full, history=history[0:-1], groupname=groupname)
                # answer = question
                user['last_response_time'] = now
                if answer is not None and len(answer) > 0:
                    history[-1] = (history[-1][0], answer)
                    user['history'] = history
                else:
                    # 如果是闲聊、或者不安全信息，删掉这句。如果仅仅是不知道答案/OCR识别的文字，留下做 context
                    if reason == 'not a question' or reason == 'security' or 'no web context':
                        del history[-1]
                    # history[-1] = (history[-1][0], '')
                    user['history'] = history

                if send and len(answer) > 1:
                    groupId = key.split('|')[-1]
                    if len(question_full) > 30:
                        format_answer = '{}..\n---\n{}'.format(question_full[0:30], answer)
                    else:
                        format_answer = '{}\n---\n{}'.format(question_full, answer)

                    black_list = ['45668581158@chatroom', '35003485913@chatroom']
                    if groupId in black_list:
                        save_sent(groupId=groupId, question_full=question_full, answer=answer, filepath='res/history_abandon.txt')
                        return None
                    save_sent(groupId=groupId, question_full=question_full, answer=answer, filepath='res/history_sent.txt')
                    send_message(groupId=groupId, answer=format_answer)

    scan_all()
    logger.debug(all_msgs)
    return None

def image_url_to_text(image_url):
    headers = {"Content-Type": "application/json"}
    text = ''
    try:
        data = {'url': image_url}
        resp = requests.post('http://{}/antiseed_ocr'.format(remote_ip_port), data=json.dumps(data), headers=headers, timeout=3)
        text = json.loads(resp.text)['text']
    except Exception as e:
        print(str(e))
    return text


def translate_doc(file_path, output_path):
    doc_en = Document(file_path)
    doc_zh = Document()
    table = doc_zh.add_table(rows=1+len(doc_en.paragraphs), cols=2)
    header_cells = table.rows[0].cells
    header_cells[0].text = '英文'
    header_cells[1].text = '中文'

    for idx,para in enumerate(doc_en.paragraphs):
        content_cells = table.rows[1+idx].cells
        content_cells[0].add_paragraph(para.text, style=para.style)

        en_text = para.text
        zh_text = ''
        if len(en_text) > 0:
            prompt = '''
请把以下体育内容翻译成中文：
{}
            '''.format(en_text)
            zh_text = llm.generate_response(prompt=prompt, remote=False)
        content_cells[1].add_paragraph(zh_text, style=para.style)
    doc_zh.save(output_path)


def process_wx_message(wx_msg):
    logger.debug(wx_msg)
    record_history(wx_msg)
    """
    普通群聊消息
    {'wcId': 'wxid_39qg5wnae8dl12', 'data': {'fromUser': 'wxid_raxq4pq3emg212', 'newMsgId': 575833538523737405, 'fromGroup': '20158567857@chatroom', 'msgId': 1615436541, 'content': '123', 'toUser': 'wxid_39qg5wnae8dl12', 'wId': '2ec3a8ba-5068-48b6-a852-0f2a40f47c4b', 'pushContent': '焕军 : 123', 'self': False, 'msgSource': '<msgsource>\n\t<silence>0</silence>\n\t<membercount>2</membercount>\n\t<signature>v1_l/B6/zGQ</sig
nature>\n\t<tmp_node>\n\t\t<publisher-id></publisher-id>\n\t</tmp_node>\n</msgsource>\n', 'timestamp': 1697340440}, 'messageType': 9, 'account': '18612393510'}
    """
    """
    自己发的消息
    {'wcId': 'wxid_39qg5wnae8dl12', 'data': {'fromUser': 'wxid_39qg5wnae8dl12', 'newMsgId': 1933322883672825931, 'fromGroup': '20158567857@chatroom', 'msgId': 1615436546, 'content': '我知道了', 'toUser': 'wxid_39qg5wnae8dl12', 'wId': '2ec3a8ba-5068-48b6-a852-0f2a40f47c4b', 'self': True, 'msgSource': '<msgsource>\n\t<signature>v1_cS75Rar5</signature>\n\t<tmp_node>\n\t\t<publisher-id></publisher-id>\n\t</tmp_node>\n</msgsource>\n', 'timestamp': 1697340874}, 'messageType': 9, 'account': '18612393510'}
    """
    """
    群里 at 消息
    {'wcId': 'wxid_39qg5wnae8dl12', 'data': {'fromUser': 'wxid_raxq4pq3emg212', 'newMsgId': 6262117166618778607, 'fromGroup': '20158567857@chatroom', 'msgId': 1615436542, 'content': '@茴香豆\u2005撤回', 'toUser': 'wxid_39qg5wnae8dl12', 'wId': '2ec3a8ba-5068-48b6-a852-0f2a40f47c4b', 'pushContent': '焕军在群聊中@了你', 'self': False, 'msgSource': '<msgsource>\n\t<atuserlist>wxid_39qg5wnae8dl12</atuserlist>\n\t<bizflag>0</bizflag>\n\t<silence>0</silence>\n\t<membercount>2</membercount>\n\t<signature>v1_f+gMVlmc</signature>\n\t<tmp_node>\n\t\t<publisher-id></publisher-id>\n\t</tmp_node>\n</msgsource>\n', 'timestamp': 1697340659}, 'messageType': 9, 'account': '18612393510'}
    """
    """
    单聊消息
     {'wcId': 'wxid_39qg5wnae8dl12', 'data': {'toUser': 'wxid_39qg5wnae8dl12', 'wId': '2ec3a8ba-5068-48b6-a852-0f2a40f47c4b', 'fromUser': 'wxid_raxq4pq3emg212', 'newMsgId': 5405961751783789412, 'msgId': 1024816412, 'self': False, 'content': '你好', 'timestamp': 1697340804}, 'messageType': 5, 'account': '18612393510'
    """
    messageType = wx_msg['messageType']
    data = wx_msg['data']
    if data['self']:
        # 自己的消息不处理
        return None

    # 群白名单
    # 测试群，测试二群，lmdeploy 2群, 贡献者群，硬件模型库，xtuner, lmdeploy 1群,
    grouplist = [
        '20158567857@chatroom', '21375247392@chatroom', '20933744776@chatroom', '21295744750@chatroom', '34744063953@chatroom', '21177113665@chatroom'
    ]

    # xtuner
    # '35003485913@chatroom', 
    # mmhuman3d 群，
    # '20814553575@chatroom', 
    grouplist += ['20814553575@chatroom']
    if os.path.exists('res/group.json'):
        with open('res/group.json') as f:
            grouplist = json.load(f)

    #  三人组
    # '11324214629@chatroom'
    # opencompass 2 群
    # grouplist += ['45668581158@chatroom']
    # 卷卷群, '18356748488@chatroom',
    # grouplist += ['18356748488@chatroom']
    # 区分是群聊普通消息，还是 at 自己的消息
    content = data['content']
    if 'fromGroup' not in data:
        print('messageType {} has no fromGroup'.format(messageType))
        return None
    groupId = data['fromGroup']

    if groupId not in grouplist:
        # 只处理某些群
        return None
    
    # 5/9/'80001' 群消息
    # 6/'80002' 群图片
    if messageType == 6 or messageType == '80002':
        # download image, parse to text

        image_url = wx_login.parse_image_url(wx_msg=wx_msg)
        if image_url is None:
            # download failed, return
            return None
        text = image_url_to_text(image_url=image_url)
        if text is None or len(text) < 1:
            return None
        
        wx_msg['data']['content'] = text
        return llm_response(wx_msg, treat_as_context=True)

    elif messageType == 5 or messageType == 9 or messageType == '80001':
        return llm_response(wx_msg)
    
    elif messageType == '80009':
        file_group_list = ['20158567857@chatroom', '21375247392@chatroom']

        if groupId not in file_group_list:
            return None
        # file
        LOCAL_EN_DOC = 'res/en.doc'
        LOCAL_ZH_DOC = 'res/zh.doc'
        code = wx_login.download_doc_file(wx_msg, file_path=LOCAL_EN_DOC)
        if code != 0:
            return None
        
        translate_doc(LOCAL_EN_DOC, LOCAL_ZH_DOC)
        doc_url, error = aliyunoss.oss_upload(LOCAL_ZH_DOC)
        if error is not None:
            print('process 60009 {}'.format(error))
            return None
        
        wx_login.send_file(wx_msg, file_url=doc_url)

    elif messageType == 14 or messageType == '80014':
        # 对于引用消息，如果要求撤回
        if 'title' in data:
            if '撤回' in data['title']:
                revert_all(groupId)
        elif '撤回' in content:
            revert_all(groupId)


def send_to_wx(response):
    headers = {"Content-Type": "application/json"}
    data = {}
    if type(response) == str:
        data = {
            'data': 'response',
        }
    else:
        data = response
    logger.debug('post {}'.format(data))
    resp = requests.post('http://{}/send'.format(remote_ip_port), data=json.dumps(data), headers=headers)
    logger.info((resp, resp.content))


def work_time():
    workTime = ['07:00:00', '22:00:00']
    dayOfWeek = datetime.datetime.now().weekday()
    #dayOfWeek = datetime.today().weekday()
    beginWork = datetime.datetime.now().strftime("%Y-%m-%d") + ' ' + workTime[0]
    endWork = datetime.datetime.now().strftime("%Y-%m-%d") + ' ' + workTime[1]
    beginWorkSeconds = time.time() - time.mktime(time.strptime(beginWork, '%Y-%m-%d %H:%M:%S'))
    endWorkSeconds = time.time() - time.mktime(time.strptime(endWork, '%Y-%m-%d %H:%M:%S'))
    if (int(dayOfWeek) in range(7)) and int(beginWorkSeconds) > 0 and int(endWorkSeconds) < 0:
        return True
    else:
        return False


if __name__ == "__main__":
    while True:
        if not work_time():
            time.sleep(60)
            continue

        # try:
        url = 'http://{}/antiseed_get'.format(remote_ip_port)
        resp = requests.post(url)
        try:
            jsonobj = json.loads(resp.text)
        except Exception as e:
            continue

        logger.debug("Received: {} ".format(resp.text))

        if len(jsonobj) < 1:
            # randval = random.randint(2, 20)
            time.sleep(2)
        else:
            process_wx_message(jsonobj)

        # except Exception as e:
        #     logger.error(str(e) + 'reconnecting')

        randval = random.randint(1, int(pow(2, 4)))
        #     time.sleep(randval)

